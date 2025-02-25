"""
 * @author: zkyuan
 * @date: 2025/2/25 11:15
 * @description:
"""
# 聚客 Ai 科技
# Kevin

# 导入 Streamlit 库，用于创建 Web 应用
import streamlit as st
# 导入递归字符文本分割器，用于将文档分割成小块
from langchain.text_splitter import RecursiveCharacterTextSplitter
# 导入 FAISS 向量存储，用于存储和检索文档嵌入
from langchain_community.vectorstores import FAISS
# 导入 OpenAI 聊天模型
from langchain_openai import ChatOpenAI
# 导入 OpenAI 嵌入模型，用于生成文本嵌入
from langchain_openai import OpenAIEmbeddings
# 导入Document类，用于封装文档内容和元数据
from langchain_core.documents import Document
# 导入对话检索链，用于处理对话和检索
from langchain.chains import ConversationalRetrievalChain
# 导入docx库，用于处理 Word 文档
import docx
# 导入 PyPDF2 库，用于处理 PDF 文档
from PyPDF2 import PdfReader

# 设置页面配置，包括标题、图标和布局
st.set_page_config(page_title="Document QA System", page_icon=":robot:", layout="wide")

# 设置页面的 CSS 样式
st.markdown(
    """<style>
.st-emotion-cache-1y6asi2 {
    height: 600px;
    background-color: #ddd;
}
.chat-message {
    padding: 10px; 
    border-radius: 10px; 
    margin-bottom: 10px;
    display: flex
}
.chat-message.user {
    # background-color: #00CC33
}
.chat-message.bot {
    # background-color: #fff
}
.chat-message .avatar {
  width: 50px;
}
.chat-message .avatar img {
  max-width: 32px;
  max-height: 32px;
  border-radius: 50%;
  object-fit: cover;
}
.chat-message .message-user {
  padding: 5px;
  border-radius: 5px;
  margin-left: 5px;
  background-color: #00FF66;
  color: #000;
}
.chat-message .message-bot {
  padding: 5px;
  border-radius: 5px;
  margin-left: 5px;
  background-color: #FFF;
  color: #000;
}
.stDeployButton {
    visibility: hidden;
}
MainMenu {visibility: hidden;}
# footer {visibility: hidden;}
# .block-container {
#     padding: 1rem 2rem;
# }
# .st-emotion-cache-16txtl3 {
#     padding: 3rem 1.5rem;
# }
</style>
# """,
    unsafe_allow_html=True,
)

# 定义机器人消息模板
bot_template = """
<div class="chat-message bot">
    <div class="avatar">
        <img src="https://www.helloimg.com/i/2024/09/12/66e28e983c1bc.png">
    </div>
    <div class="message-bot">{{MSG}}</div>
</div>
"""

# 定义用户消息模板
user_template = """
<div class="chat-message user">
    <div class="avatar">
        <img src="https://www.helloimg.com/i/2024/09/12/66e28e982c2c9.png" >
    </div>
    <div class="message-user">{{MSG}}</div>
</div>
"""

# 构建 知识索引库
# 从 PDF 文件中提取文本
def get_pdf_text(pdf_docs):
    # 存储提取的文档
    docs = []
    for document in pdf_docs:
        if document.type == "application/pdf":
            # 读取 PDF文件
            pdf_reader = PdfReader(document)
            for idx, page in enumerate(pdf_reader.pages):
                docs.append(
                    Document(
                        # 提取页面文本
                        page_content=page.extract_text(),
                        # 添加元数据
                        metadata={"source": f"{document.name} on page {idx}"},
                    )
                )
        elif document.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # 读取Word文档
            doc = docx.Document(document)
            for idx, paragraph in enumerate(doc.paragraphs):
                docs.append(
                    Document(
                        # 提取段落文本
                        page_content=paragraph.text,
                        # 添加元数据
                        metadata={"source": f"{document.name} in paragraph {idx}"},
                    )
                )
        elif document.type == "text/plain":
            # 读取纯文本文件
            text = document.getvalue().decode("utf-8")
            docs.append(Document(page_content=text, metadata={"source": document.name}))

    return docs

# 将文档分割成小块文本
def get_text_chunks(docs):
    # 创建文本分割器
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100)
    # 分割文档
    docs_chunks = text_splitter.split_documents(docs)
    return docs_chunks

# 创建向量存储
def get_vectorstore(docs_chunks):
    # 创建 OpenAI 嵌入模型
    embeddings = OpenAIEmbeddings()
    # 创建 FAISS 向量存储（基于内存、本地存储）
    vectorstore = FAISS.from_documents(docs_chunks, embedding=embeddings)
    return vectorstore

# 创建对话检索链
def get_conversation_chain(vectorstore):
    # 创建 OpenAI 聊天模型
    llm = ChatOpenAI(model="gpt-4o")
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        # 使用向量存储作为检索器
        retriever=vectorstore.as_retriever(),
        # 返回源文档
        return_source_documents=True,
    )
    return conversation_chain

# 处理用户输入并生成响应
def handle_userinput_pdf(user_question):
    # 获取聊天历史
    chat_history = st.session_state.chat_history
    # 生成响应
    response = st.session_state.conversation(
        {"question": user_question, "chat_history": chat_history}
    )
    # 添加用户问题到聊天历史
    st.session_state.chat_history.append(("user", user_question))
    # 添加机器人回答到聊天历史
    st.session_state.chat_history.append(("assistant", response["answer"]))

    # 显示用户问题
    st.write(
        user_template.replace("{{MSG}}", user_question),
        unsafe_allow_html=True,
    )

    # 获取源文档
    sources = response["source_documents"]
    # 提取源文档名称
    source_names = set([i.metadata["source"] for i in sources])
    # 合并源文档名称
    # src = "\n\n".join(source_names)
    # src = f"\n\n> source : {src}"
    message = st.session_state.chat_history[-1]
    # 显示机器人回答和源文档
    st.write(bot_template.replace("{{MSG}}", message[1]), unsafe_allow_html=True)

# 显示聊天历史记录
def show_history():
    # 获取聊天历史
    chat_history = st.session_state.chat_history

    for i, message in enumerate(chat_history):
        if i % 2 == 0:
            # 显示用户消息
            st.write(
                user_template.replace("{{MSG}}", message[1]),
                unsafe_allow_html=True,
            )
        else:
            # 显示机器人消息
            st.write(
                bot_template.replace("{{MSG}}", message[1]), unsafe_allow_html=True
            )
# 主函数
def main():
    # 显示页面标题
    st.title("Chat with documents")

    # 初始化会话状态
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    with st.sidebar:
        # 显示侧边栏标题
        st.title("Document management")
        # 文件上传控件
        pdf_docs = st.file_uploader(
            "Choose a file",
            # 支持的文件类型
            type=["pdf", "txt", "doc", "docx"],
            # 支持多文件上传
            accept_multiple_files=True,
        )
        if st.button(
                "Process documents",
                # 设置最后操作为 pdf
                on_click=lambda: setattr(st.session_state, "last_action", "pdf"),
                use_container_width=True,
        ):
            if pdf_docs:
                # 显示处理中的旋转器
                with st.spinner("Processing"):
                    # 提取 PDF文本
                    docs = get_pdf_text(pdf_docs)
                    # 分割文本
                    docs_chunks = get_text_chunks(docs)
                    # 创建向量存储
                    vectorstore = get_vectorstore(docs_chunks)
                    # 创建对话链
                    st.session_state.conversation = get_conversation_chain(vectorstore)
            else:
                # 提示用户上传文件
                st.warning("Please start by uploading the file.")

        def clear_history():
            # 清空聊天历史
            st.session_state.chat_history = []

        if st.session_state.chat_history:
            # 清空对话按钮
            st.button("Clear chat history", on_click=clear_history, use_container_width=True)

    with st.container():
        # 获取用户输入
        user_question = st.chat_input("Say something...")

    with st.container(height=400):
        # 显示聊天历史
        show_history()
        if user_question:
            if st.session_state.conversation is not None:
                # 处理用户输入
                handle_userinput_pdf(user_question)
            else:
                # 提示用户上传文件
                st.warning("Please start by uploading the file.")

    # 添加 footer
    st.markdown(
        """
        <style>
        .footer {
            # position: fixed;
            # left: 0;
            # bottom: 0;
            width: 100%;
            background-color: #f8f9fa;
            color: black;
            text-align: center;
            padding: 10px;
        }
        </style>
        <div class="footer">
        &copy;2024 聚客AI科技
        </div>
        """,
        unsafe_allow_html=True
    )

# 运行主函数
if __name__ == "__main__":
    main()